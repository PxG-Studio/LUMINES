#!/usr/bin/env python3
"""
SPARK Document Extractor
Extracts text from PDF and DOCX files for integration into comprehensive documentation
"""

import sys
import os
from pathlib import Path
from datetime import datetime

def extract_pdf_text(pdf_path):
    """Extract text from PDF file"""
    try:
        import PyPDF2
        
        text = []
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            print(f"   Found {len(reader.pages)} pages")
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    text.append(f"\n--- Page {page_num} ---\n")
                    text.append(page_text)
        
        return '\n'.join(text)
    except ImportError:
        return f"ERROR: PyPDF2 not installed. Run: pip install PyPDF2"
    except Exception as e:
        return f"ERROR extracting PDF: {str(e)}"

def extract_docx_text(docx_path):
    """Extract text from DOCX file"""
    try:
        from docx import Document
        
        doc = Document(docx_path)
        text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n'.join(text)
    except ImportError:
        return f"ERROR: python-docx not installed. Run: pip install python-docx"
    except Exception as e:
        return f"ERROR extracting DOCX: {str(e)}"

def extract_all_files():
    """Extract text from all specified files"""
    base_path = Path(r"C:\Users\nervcentre\Downloads\Indie Game Development")
    
    files_to_extract = [
        ("CM489 Capstone Proposal - Draft I.pdf", extract_pdf_text, "PDF"),
        ("Capstone Concept.docx", extract_docx_text, "DOCX"),
        ("Capstone Project Proposal.docx", extract_docx_text, "DOCX"),
    ]
    
    results = {}
    
    print("=" * 70)
    print("SPARK Document Extraction Tool")
    print("=" * 70)
    print()
    
    for filename, extractor, file_type in files_to_extract:
        file_path = base_path / filename
        
        if not file_path.exists():
            print(f"⚠️  File not found: {file_path}")
            results[filename] = None
            continue
        
        print(f"📄 Extracting {file_type}: {filename}")
        print(f"   Path: {file_path}")
        
        try:
            text = extractor(file_path)
            
            if text.startswith("ERROR"):
                print(f"   {text}")
                results[filename] = None
                continue
                
            results[filename] = text
            
            # Save to text file
            output_path = base_path / f"{Path(filename).stem}_extracted.txt"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(f"SOURCE FILE: {filename}\n")
                f.write(f"EXTRACTION DATE: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write("=" * 70 + "\n\n")
                f.write(text)
            
            print(f"   ✅ Extracted {len(text)} characters")
            print(f"   ✅ Saved to: {output_path}")
            print()
            
        except Exception as e:
            print(f"   ❌ Error: {str(e)}")
            print()
            results[filename] = None
    
    # Create combined output
    combined_path = base_path / "ALL_CAPSTONE_DOCS_EXTRACTED.txt"
    with open(combined_path, 'w', encoding='utf-8') as f:
        f.write("=" * 70 + "\n")
        f.write("SPARK CAPSTONE DOCUMENTS - COMPLETE EXTRACTION\n")
        f.write(f"Extracted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 70 + "\n\n")
        
        for filename, text in results.items():
            if text:
                f.write(f"\n{'=' * 70}\n")
                f.write(f"FILE: {filename}\n")
                f.write(f"{'=' * 70}\n\n")
                f.write(text)
                f.write("\n\n")
    
    print("=" * 70)
    print(f"✅ Combined extraction saved to: {combined_path}")
    print("=" * 70)
    
    return results

if __name__ == "__main__":
    # Check for required libraries
    missing_libs = []
    
    try:
        import PyPDF2
    except ImportError:
        missing_libs.append("PyPDF2")
    
    try:
        from docx import Document
    except ImportError:
        missing_libs.append("python-docx")
    
    if missing_libs:
        print("⚠️  Missing required libraries:")
        for lib in missing_libs:
            print(f"   - {lib}")
        print()
        print("Installing missing libraries...")
        import subprocess
        for lib in missing_libs:
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", lib])
                print(f"   ✅ Installed {lib}")
            except Exception as e:
                print(f"   ❌ Failed to install {lib}: {e}")
                print()
                print("Please install manually:")
                print(f"   pip install {' '.join(missing_libs)}")
                sys.exit(1)
        print()
    
    # Extract all files
    results = extract_all_files()
    
    # Print summary
    print("\n📊 EXTRACTION SUMMARY:")
    print("-" * 70)
    for filename, text in results.items():
        if text:
            print(f"✅ {filename}: {len(text)} characters extracted")
        else:
            print(f"❌ {filename}: Failed to extract")

